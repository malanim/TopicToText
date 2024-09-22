import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.shared import Cm
import os.path
import re
import datetime as date

year = date.date.today().year

def basic_theory(doc, plan, data):

	# Формирование стиля всего документа
	style = doc.styles['Normal']
	style.font.name = 'Times New Roman'
	style.font.size = Pt(14)

	# Содержание
	paragraph = doc.add_paragraph()
	p_fmt = paragraph.paragraph_format
	p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run('Содержание:')
	run.font.size = Pt(16)
	run.font.bold = True # Полужирный
	l = len(doc.paragraphs)
	doc.paragraphs[l-1].runs[0].add_break()
	n = 1
	for item in plan:
		doc.add_paragraph(str(n) + '. ' + item)
		n = n + 1

	# Разрыв страницы
	l = len(doc.paragraphs)
	doc.paragraphs[l-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

	n = 1
	i = 0
	for item in plan:
		paragraph = doc.add_paragraph()
		run = paragraph.add_run(str(n) + '. ' + item)
		run.font.size = Pt(16)
		run.font.bold = True # Полужирный

		paragraph = doc.add_paragraph(data[i])
		p_fmt = paragraph.paragraph_format
		p_fmt.first_line_indent = Cm(1.5)
		p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

		# Разрыв страницы
		# l = len(doc.paragraphs)
		# doc.paragraphs[l-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

		n = n + 1
		i = i + 1

	doc.add_paragraph()

	return doc

def replace_text_in_paragraph(paragraph, search, replace):
   """
   Replace occurrences of specific text in a paragraph with a new text, even when
   the text occurs across multiple consecutive runs

   Parameters:
   paragraph (docx.text.paragraph.Paragraph): The original paragraph where text 
      will be replaced.
   search (str): The text to be replaced.
   replace (str): The new text that will replace the search text.

   Returns:
   docx.text.paragraph.Paragraph: The updated paragraph with the search text
   replaced by the replace text.
   """
   # if the search text is not in the paragraph then exit
   if not re.search(search, paragraph.text): 
      return paragraph
   
   # use a join character to delimit runs, I selected the null character '\0' because it 
   # can't appear in oxml document text.
   J_CHAR = "\0"

   # join the paragraph run text with the J_CHAR as the run delimiter
   all_runs = J_CHAR.join([r.text for r in paragraph.runs])

   # compile a regex search string that allows 0,1+ occurrences of the run delimiter
   # between each character of the search string.
   pattern = re.compile(f"{J_CHAR}*".join(search))

   # subsitute the replacement text, plus the contained delimiter count in the match to 
   # keep the runs consistent
   all_runs_replaced = re.sub(pattern, 
                          lambda m: replace + m.group(0).count(J_CHAR) * J_CHAR,
                          all_runs)
   
   # iterate the paragraph runs and replace any text that has changed after the substitution
   for orig, new in zip(paragraph.runs, all_runs_replaced.split(J_CHAR)):
      if orig.text != new:
         orig.text = new
   return paragraph

def create_file_doc(tngc:list, plan:list, data:list):

	if os.path.exists('template.docx'):

		search_theme = '{{theme}}'
		search_name = '{{name}}'
		search_group = '{{group}}'
		search_colledge = '{{college}}'

		s_year = '{{year}}'

		search = [search_theme, search_name, search_group, search_colledge, s_year]

		tngc.append(year)
		
		document = Document('template.docx')
		for p in document.paragraphs:
			for i in range(len(search)):
				for l in range(len(tngc)):
					if search[i] in p.text:
						paragraph = p
						p = replace_text_in_paragraph(paragraph, str(search[i]), str(tngc[i]))
		document = basic_theory(document, plan, data)

		document.save('output.docx')

	else:

		doc = docx.Document()

		# Формирование стиля всего документа
		style = doc.styles['Normal']
		style.font.name = 'Times New Roman'
		style.font.size = Pt(14)

		# run = doc.add_paragraph().add_run()
		# r_fmt = run.font
		# r_fmt.name = 'Times New Roman'
		# r_fmt.size = Pt(14)

		# Добавление Названия учреждения по центру
		paragraph = doc.add_paragraph(tngc[3])
		p_fmt = paragraph.paragraph_format
		p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# Отступ
		for i in range(0, 10):
			l = len(doc.paragraphs)
			doc.paragraphs[l-1].runs[0].add_break()

		# Тема по центру полужирным увеличенным шрифтом
		paragraph = doc.add_paragraph()
		p_fmt = paragraph.paragraph_format
		p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
		run = paragraph.add_run('Тема работы:')
		run.font.size = Pt(16)
		run.font.bold = False # Полужирный
		paragraph = doc.add_paragraph()
		p_fmt = paragraph.paragraph_format
		p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
		run = paragraph.add_run('"' + tngc[0] + '"')
		run.font.size = Pt(18)
		run.font.bold = True # Полужирный

		# Отступ
		for i in range(0, 8):
			l = len(doc.paragraphs)
			doc.paragraphs[l-1].runs[0].add_break()

		# Добавление информации об ученике по правому краю
		paragraph = doc.add_paragraph('Подготовил(а) студент(ка)')
		p_fmt = paragraph.paragraph_format
		p_fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		paragraph = doc.add_paragraph('Группы ' + tngc[2])
		p_fmt = paragraph.paragraph_format
		p_fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		paragraph = doc.add_paragraph(tngc[1])
		p_fmt = paragraph.paragraph_format
		p_fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		# Отступ
		for i in range(0, 4):
			l = len(doc.paragraphs)
			doc.paragraphs[l-1].runs[0].add_break()

		# Год по центру
		paragraph = doc.add_paragraph(str(year), 'г.')
		p_fmt = paragraph.paragraph_format
		p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# Разрыв страницы
		l = len(doc.paragraphs)
		doc.paragraphs[l-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

		doc = basic_theory(doc, plan, data)

		doc.save('output.docx')